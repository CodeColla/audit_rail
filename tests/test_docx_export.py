"""P4-S4 — HTML to DOCX export.

Everything here asserts on the document **reopened through python-docx**, never on the raw
bytes: a .docx is a zip of XML, and a string match against it proves only that a substring
exists somewhere, not that Word will honour it.

Caveat worth knowing (`docs/phase4/01-sprint-plan.md` records it as a sprint exit criterion):
there is no OOXML validator in this environment, so "reopens cleanly in python-docx" is the
strongest automated signal available. The hand-built `w:hyperlink`, `w:numPr`, `w:tblHeader`
and `w:pBdr` elements could in principle satisfy python-docx and still be rejected by Word.
The Playwright suite renders one through docx-preview as an independent second opinion.
"""

import io
import json

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm

from api.docx_export import _STYLES, render_docx


def _doc(html: str, fmt: str = "HTML") -> Document:
    data = render_docx(title="InfoSec Policy", body_html=html, classification="INTERNAL",
                       version_label="1.0", status="PUBLISHED", content_format=fmt)
    assert data[:2] == b"PK", "a .docx is a zip; this is not one"
    return Document(io.BytesIO(data))


def _num_ids(doc) -> list[str]:
    out = []
    for p in doc.paragraphs:
        ppr = p._p.find(qn("w:pPr"))
        numpr = ppr.find(qn("w:numPr")) if ppr is not None else None
        numid = numpr.find(qn("w:numId")) if numpr is not None else None
        if numid is not None:
            out.append(numid.get(qn("w:val")))
    return out


def test_every_style_the_walker_uses_exists():
    """A style name that isn't in the template raises KeyError at render time — in front of
    a user. Pin the whole set so a python-docx upgrade fails here instead."""
    available = {s.name for s in Document().styles}
    missing = [name for name in _STYLES if name not in available]
    assert not missing, f"python-docx no longer provides: {missing}"


def test_headings_and_paragraphs():
    doc = _doc("<h1>Purpose</h1><p>Body text.</p><h2>Scope</h2>")
    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles and "Heading 2" in styles
    assert "Purpose" in [p.text for p in doc.paragraphs]


def test_nested_lists_use_indented_styles():
    doc = _doc("<ul><li><p>one</p><ul><li><p>two</p>"
               "<ul><li><p>three</p></li></ul></li></ul></li></ul>")
    styles = [p.style.name for p in doc.paragraphs]
    assert "List Bullet" in styles
    assert "List Bullet 2" in styles
    assert "List Bullet 3" in styles


def test_deeper_than_three_levels_clamps_rather_than_crashing():
    """There is no "List Bullet 4" — depth must clamp, not KeyError."""
    doc = _doc("<ul><li><ul><li><ul><li><ul><li><p>deep</p></li></ul></li></ul></li></ul></li></ul>")
    assert "deep" in [p.text for p in doc.paragraphs]


def test_sibling_ordered_lists_restart_numbering():
    """Two consecutive <ol>s must not continue 1,2,3,4. python-docx's shared numbering
    instance does exactly that, which is why each list gets a fresh w:num."""
    doc = _doc("<ol><li><p>a</p></li><li><p>b</p></li></ol><ol><li><p>fresh</p></li></ol>")
    ids = _num_ids(doc)
    assert len(set(ids)) == 2, f"expected two numbering instances, got {ids}"
    assert ids[0] == ids[1] and ids[2] != ids[0]


def test_hyperlink_round_trips():
    doc = _doc('<p>See <a href="https://rbi.org.in/x">RBI</a>.</p>')
    links = [(h.address, h.text) for p in doc.paragraphs for h in p.hyperlinks]
    assert links == [("https://rbi.org.in/x", "RBI")]


def test_unsafe_link_becomes_plain_text():
    """Defence in depth: the sanitiser already strips these, but the exporter must never
    emit a javascript: relationship even if handed one directly."""
    doc = _doc('<p><a href="javascript:alert(1)">click</a></p>')
    assert not [h for p in doc.paragraphs for h in p.hyperlinks]
    assert "click" in " ".join(p.text for p in doc.paragraphs)


def test_marks_survive():
    doc = _doc("<p><strong>b</strong><em>i</em><u>u</u><s>x</s><code>c</code></p>")
    runs = {r.text: r for p in doc.paragraphs for r in p.runs}
    assert runs["b"].bold and runs["i"].italic and runs["u"].underline
    assert runs["x"].font.strike
    assert runs["c"].font.name == "Courier New"


def test_table_is_gridded_with_a_repeating_header():
    doc = _doc("<table><tbody>"
               "<tr><th><p>Control</p></th><th><p>Owner</p></th></tr>"
               "<tr><td><p>AM-01</p></td><td><p>Ops</p></td></tr></tbody></table>")
    table = doc.tables[0]
    assert table.style.name == "Table Grid"
    assert len(table.rows) == 2 and len(table.columns) == 2
    assert table.rows[0]._tr.xml.count("tblHeader") == 1


def test_table_spans_merge():
    doc = _doc("<table><tbody>"
               "<tr><td colspan='2'><p>wide</p></td></tr>"
               "<tr><td rowspan='2'><p>tall</p></td><td><p>x</p></td></tr>"
               "<tr><td><p>y</p></td></tr></tbody></table>")
    table = doc.tables[0]
    assert len(table.rows) == 3 and len(table.columns) == 2
    assert "wide" in table.cell(0, 0).text and "tall" in table.cell(1, 0).text


def test_page_is_a4_matching_the_pdf():
    """python-docx defaults to US Letter. Compared with tolerance because page size is
    stored in twips, so Mm(210) round-trips as 7560310 EMU rather than 7560000."""
    sec = _doc("<p>x</p>").sections[0]
    assert abs(sec.page_width - Mm(210)) < 635      # < 1 twip
    assert abs(sec.page_height - Mm(297)) < 635
    assert abs(sec.left_margin - Mm(20)) < 635


# ────────────────────────────────────────────────── P5-S2: SHEET export

def test_sheet_renders_as_a_word_table():
    """content_format="SHEET" — `body_html` carries the raw JSON grid, not markup; the SHEET
    branch converts it via render.sheet_json_to_html() before the walker ever sees a tag."""
    grid = json.dumps({"data": [["Control", "Owner"], ["MFA", "Alice"]], "bold": ["A1", "B1"]})
    doc = _doc(grid, fmt="SHEET")
    table = doc.tables[0]
    assert len(table.rows) == 2 and len(table.columns) == 2
    assert table.cell(0, 0).text == "Control" and table.cell(1, 1).text == "Alice"


def test_sheet_bold_cell_is_a_bold_run_not_a_repeating_header():
    grid = json.dumps({"data": [["Control", "Owner"], ["MFA", "Alice"]], "bold": ["A1"]})
    doc = _doc(grid, fmt="SHEET")
    table = doc.tables[0]
    bold_para = table.cell(0, 0).paragraphs[0]
    assert any(r.bold for r in bold_para.runs)
    assert not any(r.bold for r in table.cell(1, 0).paragraphs[0].runs)
    # <strong>, not <th> — a bold first row must not become a page-repeating header
    assert table.rows[0]._tr.xml.count("tblHeader") == 0


def test_sheet_cell_alignment_is_applied():
    grid = json.dumps({"data": [["x", "y"]], "align": {"B1": "right"}})
    doc = _doc(grid, fmt="SHEET")
    table = doc.tables[0]
    assert table.cell(0, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.cell(0, 0).paragraphs[0].alignment is None


def test_markdown_versions_still_export():
    """Pre-S4 versions hold markdown. They must come out as real headings, not as a literal
    "# Purpose"."""
    doc = _doc("# Purpose\n\nOld **markdown** body.", fmt="MARKDOWN")
    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles
    assert any("Purpose" == p.text for p in doc.paragraphs)


def test_code_block_keeps_its_lines():
    doc = _doc('<pre><code class="language-python">a = 1\nb = 2</code></pre>')
    macro = [p for p in doc.paragraphs if p.style.name == "macro"]
    assert macro and "a = 1" in macro[0].text


@pytest.mark.parametrize("html", [
    "",                                              # empty document
    "bare text with no element wrapper",
    "<p></p>",                                       # empty paragraph
    "<p>caf&eacute; &mdash; ✓ 中文</p>",  # entities + unicode
    "<blockquote>no inner block</blockquote>",
    "<ul><li>bare li, no p</li></ul>",               # markdown-shaped list item
    "<li>orphan list item</li>",
    "<table></table>",                               # table with no rows
    "<table><tbody><tr></tr></tbody></table>",
    "<p><strong><em><u>triple nested</u></em></strong></p>",
    "<ul><li><p>item</p><table><tbody><tr><td>t</td></tr></tbody></table></li></ul>",
    '<p><img src="https://evil.test/x.png" alt="diagram"></p>',
    "<h1>a</h1>" * 200,                              # long document
])
def test_pathological_input_does_not_raise(html):
    """A malformed or unusual body must never 500 the export endpoint."""
    assert _doc(html) is not None


def test_image_becomes_a_placeholder_and_is_never_fetched():
    """If an <img> ever reaches the exporter (legacy markdown can contain one), it must
    render as text. Fetching it would be the same SSRF the sanitiser exists to prevent."""
    doc = _doc('<p><img src="https://evil.test/x.png" alt="network diagram"></p>')
    assert "network diagram" in " ".join(p.text for p in doc.paragraphs)
