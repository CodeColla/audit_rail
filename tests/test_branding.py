"""P6-S5 — the letterhead an exported document actually carries.

This is a multi-tenant correctness fix wearing a feature's clothes. Before P6-S5,
`render.render_pdf` had no `org` parameter at all and `routers/documents.py` never read
`tenants.name`, so **every customer's approved, signed policy was letterheaded
"KIAM INTL PVT LTD"** — the first customer's company. Nothing failed; it just quietly said the
wrong thing on the most formal artefact the product produces.

The tests below are therefore mostly about the boring cases — no logo, a logo whose bytes have
gone missing, another tenant's row — because those are the ones that would make an export
either lie or 500, and an export is the last place a customer will tolerate either.
"""

import io
import re
import uuid

from sqlalchemy import text as sqltext

#: A real one-pixel PNG, not just the magic bytes. The upload sniffer only checks the first
#: eight bytes, but python-docx decodes the whole image — so a fake PNG passes the API and
#: then silently fails `add_picture`, which is exactly the fail-soft path and therefore
#: exactly the wrong fixture for proving the logo arrives.
PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d4948445200000001000000010802000000907753de"
    "0000000c49444154789c63e095d20600008a005328b499f20000000049454e44"
    "ae426082")


def _new_org(client, tag):
    r = client.post("/api/auth/signup", json={
        "full_name": f"{tag} owner", "email": f"{tag}@example.com",
        "password": "Passw0rdOne", "organisation_name": f"Org {tag} Pvt Ltd"})
    assert r.status_code == 201, r.text
    j = r.json()
    return j, {"Authorization": f"Bearer {j['access_token']}"}


def _letterhead(tenant_id):
    from api.rendering import branding
    from api.core.database import engine

    with engine.connect() as conn:
        return branding.letterhead(conn, tenant_id)


# ────────────────────────────────────────────── the lookup itself

def test_the_tenants_own_name_is_what_comes_back(app_client):
    tag = f"brand{uuid.uuid4().hex[:6]}"
    org, _h = _new_org(app_client, tag)
    brand = _letterhead(org["tenant_id"])
    assert brand["org"] == f"Org {tag} Pvt Ltd"
    assert brand["org"] != "KIAM INTL PVT LTD"


def test_no_logo_is_not_an_error(app_client):
    """Most organisations never upload one. `letterhead` must answer "no logo" rather than
    raise, or the export path breaks for the majority of customers."""
    org, _h = _new_org(app_client, f"nologo{uuid.uuid4().hex[:6]}")
    brand = _letterhead(org["tenant_id"])
    assert brand["logo"] is None and brand["logo_data_uri"] is None
    assert brand["org"]                      # …but the name is still real


def test_an_uploaded_logo_comes_back_in_both_forms(app_client):
    """Two forms because the two exports need different ones: a `data:` URI for xhtml2pdf
    (which decodes it with no filesystem and no network) and raw bytes for python-docx's
    `add_picture`."""
    org, h = _new_org(app_client, f"logo{uuid.uuid4().hex[:6]}")
    app_client.put("/api/org/logo", headers=h, files={"file": ("m.png", PNG, "image/png")})

    brand = _letterhead(org["tenant_id"])
    assert brand["logo"] == PNG
    assert brand["logo_mime"] == "image/png"
    assert brand["logo_data_uri"].startswith("data:image/png;base64,")


def test_a_logo_whose_bytes_have_vanished_degrades_instead_of_raising(app_client):
    """A vault object can go missing — a restore, a botched migration, a full disk. Publishing
    a policy is not the moment to discover it, and a document with no logo is still a valid
    document. The alternative is a 500 on export for every document in the tenant."""
    from api.core import storage
    from api.core.database import engine

    org, h = _new_org(app_client, f"gone{uuid.uuid4().hex[:6]}")
    app_client.put("/api/org/logo", headers=h, files={"file": ("m.png", PNG, "image/png")})
    with engine.connect() as conn:
        key = conn.execute(sqltext(
            "SELECT f.storage_key FROM tenants t JOIN files f ON f.id = t.logo_file_id "
            "WHERE t.id = :t"), {"t": org["tenant_id"]}).scalar()
    storage.path_for(key).unlink()

    brand = _letterhead(org["tenant_id"])
    assert brand["logo"] is None and brand["logo_data_uri"] is None
    assert brand["org"].endswith("Pvt Ltd")


def test_an_unknown_tenant_falls_back_rather_than_crashing():
    brand = _letterhead(str(uuid.uuid4()))
    assert brand["org"] == "KIAM INTL PVT LTD"      # the last-resort default, by design
    assert brand["logo"] is None


# ────────────────────────────────────────────── end to end, through the export routes

def _a_published_document(client, h):
    """A tenant's own document, published, so both export routes are reachable."""
    from api.core.database import engine

    me = client.get("/api/auth/me", headers=h).json()
    tid = me.get("tenant_id") or me["organisations"][0]["tenant_id"]
    pid = str(uuid.uuid4())
    with engine.begin() as c:
        c.execute(sqltext("INSERT INTO people (id,tenant_id,full_name,email) "
                          "VALUES (:i,:t,'Owner',:e)"),
                  {"i": pid, "t": tid, "e": f"o-{uuid.uuid4().hex[:6]}@x.example"})
    r = client.post("/api/documents", headers=h, json={
        "title": "Information Security Policy", "owner_person_id": pid,
        "document_type": "POLICY", "content": "<p>Body</p>", "content_format": "HTML"})
    assert r.status_code == 201, r.text
    return r.json()["id"], r.json()["version_id"]


def test_both_exports_carry_the_callers_own_organisation(app_client):
    """The end of the defect: a draft PDF and a Word export must say who they belong to."""
    org, h = _new_org(app_client, f"e2e{uuid.uuid4().hex[:6]}")
    doc_id, ver_id = _a_published_document(app_client, h)
    name = _letterhead(org["tenant_id"])["org"]

    pdf = app_client.get(f"/api/documents/{doc_id}/versions/{ver_id}/render.pdf", headers=h)
    assert pdf.status_code == 200 and pdf.content[:4] == b"%PDF"
    from pypdf import PdfReader
    text = re.sub(r"\s+", " ", PdfReader(io.BytesIO(pdf.content)).pages[0].extract_text() or "")
    assert name in text
    assert "KIAM INTL PVT LTD" not in text

    docx = app_client.get(f"/api/documents/{doc_id}/versions/{ver_id}/render.docx", headers=h)
    assert docx.status_code == 200
    from docx import Document
    header = Document(io.BytesIO(docx.content)).sections[0].header.paragraphs[0].text
    assert name in header
    assert "KIAM INTL PVT LTD" not in header


def test_the_logo_reaches_the_word_header_as_a_picture(app_client):
    """python-docx puts a picture in a header only if it is added as a RUN. Assigning
    `paragraph.text` — which is what this code did before the logo existed — replaces every
    run in the paragraph and would silently delete it."""
    org, h = _new_org(app_client, f"wlogo{uuid.uuid4().hex[:6]}")
    app_client.put("/api/org/logo", headers=h, files={"file": ("m.png", PNG, "image/png")})
    doc_id, ver_id = _a_published_document(app_client, h)

    docx = app_client.get(f"/api/documents/{doc_id}/versions/{ver_id}/render.docx", headers=h)
    assert docx.status_code == 200
    from docx import Document
    header = Document(io.BytesIO(docx.content)).sections[0].header
    xml = header.paragraphs[0]._p.xml
    assert "graphicData" in xml or "blip" in xml, "no picture run in the section header"
    # …and the text is still there beside it
    assert _letterhead(org["tenant_id"])["org"] in header.paragraphs[0].text


def test_a_corrupt_logo_never_fails_an_export(app_client):
    """Fail-soft, deliberately: a policy export must not 500 because the logo was odd. The
    magic-byte sniffer makes this hard to reach through the API, so it is forced here."""
    from api.rendering import docx_export, render

    meta_args = dict(title="T", body_html="<p>x</p>", classification="INTERNAL",
                     version_label="1.0", org="Acme Ltd")
    data = docx_export.render_docx(**meta_args, logo=b"\x89PNG\r\n\x1a\nnot really a png",
                                   logo_mime="image/png")
    from docx import Document
    assert "Acme Ltd" in Document(io.BytesIO(data)).sections[0].header.paragraphs[0].text

    pdf, engine = render.render_pdf(title="T", body_md="<p>x</p>", classification="INTERNAL",
                                    version_label="1.0", content_format="HTML",
                                    org="Acme Ltd", logo_data_uri="data:image/png;base64,!!!")
    assert pdf[:4] == b"%PDF"
    assert engine == "xhtml2pdf", "a bad logo must not knock the renderer down to the stub"
