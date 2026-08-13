"""DOCX export for authored documents (P4-S4).

Banks ask for policies as Word files, so a version has to leave as .docx as well as .pdf.

Why a hand-written walker rather than a converter library: the HTML here is not arbitrary web
content. It is the output of a TipTap editor that has already passed through
`api.html_sanitize`, so the tag set is a fixed, known list of about twenty. A recursive walk
over that list is smaller, dependency-free (python-docx and lxml are already installed) and —
crucially — testable tag by tag. The one plausible library, htmldocx, was last released in
2021 and would add beautifulsoup4.

Two OOXML details that python-docx does not wrap, and which are therefore built by hand here:

* **Numbering.** `add_paragraph(style="List Bullet")` reuses one shared numbering instance, so
  two consecutive ordered lists continue 1,2,3,4 instead of restarting. Each top-level list
  therefore gets a fresh `w:num` pointing at the style's own `w:abstractNumId`, derived at
  runtime rather than hardcoded so it survives a python-docx upgrade.
* **Hyperlinks.** There is no `add_hyperlink`; the `w:hyperlink` element and its relationship
  are assembled directly.

The layout mirrors `api.render`'s PDF — same A4 page, same letterhead and footer via
`render.doc_meta` — so the two exports of one version look like the same document.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Mm, Pt, RGBColor
from lxml import html as lhtml

from api.rendering import imagefile, render
from api.rendering.html_sanitize import IMG_SRC_RE

#: Style names verified present on a default python-docx Document(); `tests/test_docx_export.py`
#: re-checks this so an upgrade that renames one fails loudly instead of raising KeyError in
#: front of a user. List depth clamps at 3 — "List Bullet 4" does not exist.
_STYLES = (
    "Normal", "Title", "Caption", "Quote", "macro", "Table Grid", "List Continue",
    *(f"Heading {i}" for i in range(1, 7)),
    *(f"List Bullet{s}" for s in ("", " 2", " 3")),
    *(f"List Number{s}" for s in ("", " 2", " 3")),
)

_ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
          "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
_SAFE_SCHEMES = ("http://", "https://", "mailto:")
_MARK_TAGS = {"strong": "bold", "b": "bold", "em": "italic", "i": "italic",
              "u": "underline", "s": "strike", "del": "strike", "strike": "strike"}


# ── numbering plumbing ──────────────────────────────────────────────────────────────────
def _abstract_for(doc, style_name: str) -> str | None:
    """The abstractNumId behind a list style, resolved through numbering.xml."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None
    ppr = style.element.find(qn("w:pPr"))
    numpr = ppr.find(qn("w:numPr")) if ppr is not None else None
    numid = numpr.find(qn("w:numId")) if numpr is not None else None
    if numid is None:
        return None
    want = numid.get(qn("w:val"))
    for num in doc.part.numbering_part.element.findall(qn("w:num")):
        if num.get(qn("w:numId")) == want:
            abs_el = num.find(qn("w:abstractNumId"))
            return abs_el.get(qn("w:val")) if abs_el is not None else None
    return None


def _fresh_num(doc, abstract_id: str) -> int:
    """A brand-new numbering instance, so this list restarts at 1."""
    numbering = doc.part.numbering_part.element
    used = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    new_id = (max(used) if used else 0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abs_el = OxmlElement("w:abstractNumId")
    abs_el.set(qn("w:val"), abstract_id)
    num.append(abs_el)
    numbering.append(num)
    return new_id


def _set_numpr(par, num_id: int, level: int = 0) -> None:
    ppr = par._p.get_or_add_pPr()
    for old in ppr.findall(qn("w:numPr")):
        ppr.remove(old)
    numpr = OxmlElement("w:numPr")
    for tag, val in (("w:ilvl", str(level)), ("w:numId", str(num_id))):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        numpr.append(el)
    ppr.append(numpr)


def _text(el) -> str:
    return " ".join((el.text_content() or "").split())


class _Walker:
    """Recursive HTML -> python-docx emitter over the sanitiser's tag set."""

    def __init__(self, doc, image_resolver=None):
        self.doc = doc
        self.image_resolver = image_resolver
        self._abstract: dict[str, str | None] = {}

    # ---- inline ----
    def _hyperlink(self, par, url: str, text: str, marks: dict) -> None:
        if not url.lower().startswith(_SAFE_SCHEMES):
            self._run(par, text, marks)          # unsafe scheme -> plain text, never a link
            return
        rid = par.part.relate_to(url, RT.HYPERLINK, is_external=True)
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), rid)
        run = par.add_run(text)
        run.underline = True
        link.append(run._r)
        par._p.append(link)

    def _run(self, par, text: str, marks: dict):
        if not text:
            return None
        run = par.add_run(text)
        run.bold = marks.get("bold") or None
        run.italic = marks.get("italic") or None
        run.underline = marks.get("underline") or None
        if marks.get("strike"):
            run.font.strike = True
        if marks.get("code"):
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        return run

    def _inline(self, par, el, marks: dict) -> None:
        """Walk an element's inline children, accumulating character marks."""
        self._run(par, el.text or "", marks)
        for child in el:
            tag = child.tag if isinstance(child.tag, str) else ""
            if tag == "br":
                par.add_run().add_break()
            elif tag == "a":
                self._hyperlink(par, child.get("href") or "", child.text_content() or "", marks)
            elif tag in _MARK_TAGS:
                self._inline(par, child, {**marks, _MARK_TAGS[tag]: True})
            elif tag == "code":
                self._inline(par, child, {**marks, "code": True})
            elif tag == "img":
                if not self._picture(par, child):
                    self._run(par, f"[image omitted: {child.get('alt') or 'image'}]", marks)
            else:
                self._inline(par, child, marks)
            self._run(par, child.tail or "", marks)

    #: The usable text width from `_page_setup`: A4 minus 20mm each side.
    _MAX_PICTURE_MM = 170

    def _picture(self, par, el) -> bool:
        """Embed one image, or return False so the caller writes the placeholder instead.

        **Nothing is ever fetched.** The src is matched against `html_sanitize.IMG_SRC_RE` and
        the id handed to the injected resolver, which reads bytes we already hold, scoped to
        the tenant. A remote URL never matches, so the pre-P6-S5 promise — that a Word export
        cannot be turned into a server-side fetch — is unchanged.

        Fails soft on anything odd: python-docx raises for formats it cannot parse (it has no
        WebP support at all, hence `to_docx_safe`) and for corrupt bytes, and one strange
        picture must not cost the author the whole export.
        """
        if self.image_resolver is None:
            return False
        found = IMG_SRC_RE.match(el.get("src") or "")
        if not found:
            return False
        got = self.image_resolver(found.group(1))
        if not got:
            return False
        data, mime = got
        try:
            width_px = int(el.get("width") or 0)
        except ValueError:
            width_px = 0
        # CSS pixels are 1/96in; Word wants a real measure. Clamped to the text column so a
        # 4000px screenshot cannot run off the page.
        width_mm = min(width_px * 25.4 / 96, self._MAX_PICTURE_MM) if width_px else None
        try:
            run = par.add_run()
            if width_mm:
                run.add_picture(io.BytesIO(imagefile.to_docx_safe(data, mime)),
                                width=Mm(width_mm))
            else:
                run.add_picture(io.BytesIO(imagefile.to_docx_safe(data, mime)))
            return True
        except Exception:
            return False

    def _para(self, el, style: str = "Normal", marks: dict | None = None):
        par = self.doc.add_paragraph(style=style)
        align = (el.get("style") or "").replace(" ", "")
        for key, val in _ALIGN.items():
            if f"text-align:{key}" in align:
                par.alignment = val
        self._inline(par, el, marks or {})
        return par

    # ---- blocks ----
    def _list(self, el, depth: int = 0) -> None:
        ordered = el.tag == "ol"
        lvl = min(depth, 2)
        style = f"List {'Number' if ordered else 'Bullet'}{'' if lvl == 0 else f' {lvl + 1}'}"
        if style not in self._abstract:
            self._abstract[style] = _abstract_for(self.doc, style)
        abstract = self._abstract[style]
        num_id = _fresh_num(self.doc, abstract) if abstract else None

        for li in el.findall("li"):
            blocks = [c for c in li if isinstance(c.tag, str)]
            nested = [c for c in blocks if c.tag in ("ul", "ol")]
            own = [c for c in blocks if c.tag not in ("ul", "ol")]
            first = True

            # a bare <li>text</li> and a TipTap <li><p>text</p></li> must render alike
            if (li.text or "").strip() or not own:
                par = self.doc.add_paragraph(style=style)
                if num_id:
                    _set_numpr(par, num_id, lvl)
                self._inline(par, li, {})
                first = False

            for blk in own:
                if first:
                    par = self._para(blk, style=style)
                    if num_id:
                        _set_numpr(par, num_id, lvl)
                    first = False
                else:
                    self._para(blk, style="List Continue")

            for sub in nested:
                self._list(sub, depth + 1)

    def _rule(self) -> None:
        par = self.doc.add_paragraph()
        ppr = par._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "auto")):
            bottom.set(qn(k), v)
        borders.append(bottom)
        ppr.append(borders)

    @staticmethod
    def _cell_style(cell, par, css: str) -> None:
        """Apply a spreadsheet cell's `style=""` to a Word cell.

        `css` arrives whitespace-stripped. Values were validated server-side by
        `render._STYLE_PROPS` before ever being stored, so the parsing here can stay simple —
        but it still fails soft: a declaration it doesn't recognise is skipped, never raised,
        because losing one cell's colour must not fail a whole document export.
        """
        for key, val in _ALIGN.items():
            if f"text-align:{key}" in css:
                par.alignment = val

        m = re.search(r"font-size:(\d+(?:\.\d+)?)pt", css)
        if m:
            for run in par.runs:
                run.font.size = Pt(float(m.group(1)))

        m = re.search(r"(?<!-)color:#([0-9a-fA-F]{3}|[0-9a-fA-F]{6})\b", css)
        if m:
            hexv = m.group(1)
            if len(hexv) == 3:                      # #abc -> #aabbcc
                hexv = "".join(ch * 2 for ch in hexv)
            for run in par.runs:
                run.font.color.rgb = RGBColor.from_string(hexv.upper())

        # Cell shading is not wrapped by python-docx — build the w:shd element by hand,
        # same approach as the w:tblHeader/w:numPr elements elsewhere in this module.
        m = re.search(r"background-color:#([0-9a-fA-F]{3}|[0-9a-fA-F]{6})\b", css)
        if m:
            hexv = m.group(1)
            if len(hexv) == 3:
                hexv = "".join(ch * 2 for ch in hexv)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hexv.upper())
            cell._tc.get_or_add_tcPr().append(shd)

    def _table(self, el) -> None:
        rows = [r for r in el.iter("tr")]
        if not rows:
            return
        grid = [[c for c in r if isinstance(c.tag, str) and c.tag in ("td", "th")] for r in rows]
        ncols = max((sum(int(c.get("colspan", 1) or 1) for c in r) for r in grid), default=1)
        table = self.doc.add_table(rows=len(rows), cols=ncols)
        table.style = self.doc.styles["Table Grid"]

        taken: set[tuple[int, int]] = set()
        for ri, cells in enumerate(grid):
            ci = 0
            for cell in cells:
                while (ri, ci) in taken:
                    ci += 1
                if ci >= ncols:
                    break
                cspan = max(1, min(int(cell.get("colspan", 1) or 1), ncols - ci))
                rspan = max(1, min(int(cell.get("rowspan", 1) or 1), len(rows) - ri))
                target = table.cell(ri, ci)
                if cspan > 1 or rspan > 1:
                    target = target.merge(table.cell(ri + rspan - 1, ci + cspan - 1))
                for r_ in range(ri, ri + rspan):
                    for c_ in range(ci, ci + cspan):
                        taken.add((r_, c_))
                par = target.paragraphs[0]
                self._inline(par, cell, {"bold": True} if cell.tag == "th" else {})
                # Cell-level style on a <td>/<th>. Cells need this separately from `_para`
                # (which does the same for a block-level <p>) because they never go through
                # it. The set here must stay in step with render._STYLE_PROPS — anything a
                # spreadsheet can store and the PDF can show has to survive into Word too,
                # or formatting silently disappears in one export but not the other.
                self._cell_style(target, par, (cell.get("style") or "").replace(" ", ""))
                ci += cspan

        # repeat an all-<th> first row as a header on every page
        if grid and grid[0] and all(c.tag == "th" for c in grid[0]):
            trpr = rows and table.rows[0]._tr.get_or_add_trPr()
            trpr.append(OxmlElement("w:tblHeader"))

    def block(self, el) -> None:
        tag = el.tag if isinstance(el.tag, str) else ""
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._para(el, style=f"Heading {tag[1]}")
        elif tag == "p":
            if _text(el) or len(el):
                self._para(el)
        elif tag in ("ul", "ol"):
            self._list(el)
        elif tag == "blockquote":
            children = [c for c in el if isinstance(c.tag, str)]
            if children:
                for child in children:
                    self._para(child, style="Quote")
            else:
                self._para(el, style="Quote")
        elif tag == "pre":
            par = self.doc.add_paragraph(style="macro")
            par.add_run(el.text_content() or "")     # python-docx turns \n into <w:br/>
        elif tag == "hr":
            self._rule()
        elif tag == "table":
            self._table(el)
        elif tag in ("div", "section", "article", "colgroup"):
            for child in el:
                if isinstance(child.tag, str):
                    self.block(child)
        elif tag:
            self._para(el)


def _page_setup(doc, meta: dict, logo: bytes | None = None,
                logo_mime: str | None = None) -> None:
    """A4 with the PDF's margins — python-docx defaults to US Letter."""
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin = sec.right_margin = Mm(20)
    sec.top_margin = sec.bottom_margin = Mm(22)

    head = sec.header.paragraphs[0]
    # Runs, not `head.text = …`. Assigning `.text` REPLACES every run in the paragraph, so it
    # would silently delete the logo picture added just before it — the kind of change that
    # looks harmless in a diff and produces a Word file with no logo.
    if logo:
        try:
            head.add_run().add_picture(
                io.BytesIO(imagefile.to_docx_safe(logo, logo_mime)), height=Mm(7))
            head.add_run("  ")
        except Exception:
            # A bad logo must never fail a policy export. Same fail-soft discipline the rest
            # of this module applies to odd markup: degrade to the text-only header.
            pass
    head.add_run(f"{meta['org']}\t{meta['classification']}")
    foot = sec.footer.paragraphs[0]
    foot.text = meta["footer_line"]
    for par in (head, foot):
        for run in par.runs:
            run.font.size = Pt(7.5)
        # the built-in Header/Footer styles hard-code Letter's tab stops; re-place them
        par.paragraph_format.tab_stops.clear_all()
        usable = sec.page_width - sec.left_margin - sec.right_margin
        par.paragraph_format.tab_stops.add_tab_stop(usable)


def render_docx(*, title: str, body_html: str, classification: str, version_label: str,
                org: str = render.DEFAULT_ORG, status: str = "DRAFT",
                content_format: str = "MARKDOWN",
                logo: bytes | None = None, logo_mime: str | None = None,
                image_resolver=None) -> bytes:
    """Render one document version as a .docx. Always on the fly; nothing is persisted."""
    # pre-S4 versions hold markdown; convert so they export as headings rather than "# x"
    html = body_html or ""
    if content_format == "SHEET":
        # The one conversion a spreadsheet needs — from here it's an HTML <table>, which
        # _Walker._table() (below) already knows how to turn into a Word table.
        html = render.sheet_json_to_html(html)
    elif content_format != "HTML":
        html = render.md_to_html(html)

    meta = render.doc_meta(title=title, classification=classification,
                           version_label=version_label, org=org, status=status)
    doc = Document()
    _page_setup(doc, meta, logo, logo_mime)

    doc.add_paragraph(meta["header_line"], style="Caption")
    doc.add_paragraph(title, style="Title")

    walker = _Walker(doc, image_resolver)
    root = lhtml.fragment_fromstring(html, create_parent="div")
    for el in root:
        if isinstance(el.tag, str):
            walker.block(el)
    if not len(root) and (root.text or "").strip():
        doc.add_paragraph(root.text.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
